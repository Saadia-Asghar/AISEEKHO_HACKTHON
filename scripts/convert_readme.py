import os
import re

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    import subprocess
    import sys
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_element(name):
    return OxmlElement(name)

def set_cell_border(cell, **kwargs):
    """
    kwargs can be: top, bottom, left, right, insideH, insideV
    value is a dictionary like: {'sz': 12, 'val': 'single', 'color': 'FF0000', 'space': '0'}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def convert_md_to_docx(md_path, docx_path):
    print(f"Reading {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    print("Converting content...")
    in_code_block = False
    in_table = False
    table_headers = []
    table_rows = []

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("🟢 KhidmatAI — Bolein, Hum Karein")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED) # Violet Brand Color
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("AI-Powered Service Orchestrator for Pakistan’s Informal Economy\nGoogle Antigravity Hackathon 2026 — Challenge 2 Submission")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x7B, 0x74, 0x87)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip original title line since we added it customized above
        if line.startswith("# 🟢 KhidmatAI"):
            i += 1
            continue
        if line.startswith("**AI-Powered Service Orchestrator"):
            i += 1
            continue
        if line.startswith("*Google Antigravity Hackathon"):
            i += 1
            continue

        # Code Block
        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            # Add simple visual background tint simulating code highlights
            i += 1
            continue

        # Table Parsing
        if "|" in line:
            if "---" in line: # Divider
                i += 1
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                # Write gathered table
                rows_count = len(table_rows) + 1
                cols_count = max(len(table_headers), max(len(r) for r in table_rows) if table_rows else 0)
                table = doc.add_table(rows=rows_count, cols=cols_count)
                table.autofit = True
                
                # Format Header
                hdr_cells = table.rows[0].cells
                for col_idx, text in enumerate(table_headers):
                    if col_idx < len(hdr_cells):
                        hdr_cells[col_idx].text = text
                        set_cell_shading(hdr_cells[col_idx], "7C3AED") # Violet Brand Header
                        set_cell_margins(hdr_cells[col_idx])
                        run = hdr_cells[col_idx].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                
                # Format Data Rows
                for row_idx, r_data in enumerate(table_rows):
                    row_cells = table.rows[row_idx + 1].cells
                    bg_color = "F9F8FD" if row_idx % 2 == 0 else "FFFFFF"
                    for col_idx, text in enumerate(r_data):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = text
                            set_cell_shading(row_cells[col_idx], bg_color)
                            set_cell_margins(row_cells[col_idx])
                            # Set light grid borders
                            set_cell_border(row_cells[col_idx], bottom={'sz': 2, 'val': 'single', 'color': 'E1E2E4'})
                
                doc.add_paragraph() # Spacer
                in_table = False
                table_headers = []
                table_rows = []

        # Headings
        if line.startswith("## "):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(line[3:])
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
            i += 1
            continue
        elif line.startswith("### "):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(line[4:])
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x5a, 0x00, 0xc6)
            i += 1
            continue

        # Bullet lists
        if line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            # Handle italic/bold in bullets
            cleaned_text = line[2:]
            # Simple bold/italic strip for clean word doc rendering
            cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned_text)
            cleaned_text = re.sub(r'\*(.*?)\*', r'\1', cleaned_text)
            p.add_run(cleaned_text)
            i += 1
            continue

        # Normal text paragraph
        if line:
            p = doc.add_paragraph()
            # Simple bold/italic regex formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph()

        i += 1

    print(f"Saving to {docx_path}...")
    doc.save(docx_path)
    print("Successfully converted README.md to README.docx!")

if __name__ == "__main__":
    convert_md_to_docx("README.md", "README.docx")
